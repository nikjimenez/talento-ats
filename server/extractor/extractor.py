"""
extractor.py — servicio de extracción de datos de hojas de vida.

Por qué Python y no Node: el ecosistema de análisis de documentos vive
aquí. `pdfplumber` lee PDF con posiciones reales, `python-docx` abre Word,
y spaCy hace reconocimiento de entidades en español. En Node habría que
reimplementar todo eso a mano.

Qué NO hace este servicio: no toca la base de datos, no valida sesiones,
no guarda archivos. Recibe bytes, devuelve datos. Toda la persistencia y
la autorización siguen en el servidor Node, que es el único que habla con
PostgreSQL.

Arranque:  uvicorn extractor:app --host 127.0.0.1 --port 8100
"""

from __future__ import annotations

import io
import re
import unicodedata
from dataclasses import dataclass, field, asdict
from datetime import date

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse

app = FastAPI(
    title="Talento ATS · Extractor de hojas de vida",
    version="1.0.0",
    docs_url="/docs",
)

MAX_BYTES = 10 * 1024 * 1024


# ─────────────────────────── Utilidades de texto ───────────────────────────

def sin_tildes(texto: str) -> str:
    """Normaliza para comparar: 'Bogotá' y 'Bogota' deben coincidir."""
    return "".join(
        c for c in unicodedata.normalize("NFD", texto)
        if unicodedata.category(c) != "Mn"
    )


def limpiar(texto: str) -> str:
    texto = texto.replace("\xa0", " ")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


# ─────────────────────────── Lectura de archivos ───────────────────────────

def leer_pdf(datos: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise HTTPException(500, "Falta pdfplumber. Instala las dependencias.")

    partes: list[str] = []
    with pdfplumber.open(io.BytesIO(datos)) as pdf:
        for pagina in pdf.pages[:15]:          # tope: una HV no tiene 40 páginas
            partes.append(pagina.extract_text() or "")
    return limpiar("\n".join(partes))


def leer_docx(datos: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise HTTPException(500, "Falta python-docx. Instala las dependencias.")

    documento = docx.Document(io.BytesIO(datos))
    partes = [p.text for p in documento.paragraphs]
    for tabla in documento.tables:
        for fila in tabla.rows:
            partes.append(" ".join(c.text for c in fila.cells))
    return limpiar("\n".join(partes))


def extraer_texto(datos: bytes, nombre: str) -> str:
    if datos[:4] == b"%PDF":
        return leer_pdf(datos)
    if datos[:4] == b"PK\x03\x04":              # docx es un zip
        return leer_docx(datos)
    if nombre.lower().endswith(".txt"):
        return limpiar(datos.decode("utf-8", errors="replace"))
    raise HTTPException(415, "Formato no soportado. Se aceptan PDF, DOCX y TXT.")


# ─────────────────────────── Patrones colombianos ──────────────────────────

# Cédula: 6 a 10 dígitos, con o sin puntos, precedida de una etiqueta.
RE_CEDULA = re.compile(
    r"(?:c\.?\s*c\.?|cedula|cédula|documento|identificacion|identificación|nit)"
    r"[\s.:#nroº°-]*([\d][\d.,\s]{5,14}\d)",
    re.IGNORECASE,
)

# Celular colombiano: empieza en 3, diez dígitos. Acepta +57 y separadores.
RE_CELULAR = re.compile(
    r"(?:\+?57[\s.-]?)?(3\d{2})[\s.-]?(\d{3})[\s.-]?(\d{4})\b"
)

RE_FIJO = re.compile(r"(?:\+?57[\s.-]?)?\(?([1-8])\)?[\s.-]?(\d{3})[\s.-]?(\d{4})\b")

RE_EMAIL = re.compile(r"\b[\w.+-]+@[\w-]+\.[\w.-]{2,}\b")

RE_FECHA = re.compile(
    r"\b(\d{1,2})[/\-\s]+(?:de\s+)?"
    r"(ene|feb|mar|abr|may|jun|jul|ago|sep|set|oct|nov|dic|\d{1,2})"
    r"[a-z]*[/\-\s]+(?:de\s+)?(\d{4})\b",
    re.IGNORECASE,
)

MESES = {
    "ene": 1, "feb": 2, "mar": 3, "abr": 4, "may": 5, "jun": 6,
    "jul": 7, "ago": 8, "sep": 9, "set": 9, "oct": 10, "nov": 11, "dic": 12,
}

# Los 32 departamentos más Bogotá D.C.
DEPARTAMENTOS = [
    "Amazonas", "Antioquia", "Arauca", "Atlántico", "Bolívar", "Boyacá",
    "Caldas", "Caquetá", "Casanare", "Cauca", "Cesar", "Chocó", "Córdoba",
    "Cundinamarca", "Guainía", "Guaviare", "Huila", "La Guajira", "Magdalena",
    "Meta", "Nariño", "Norte de Santander", "Putumayo", "Quindío", "Risaralda",
    "San Andrés y Providencia", "Santander", "Sucre", "Tolima",
    "Valle del Cauca", "Vaupés", "Vichada", "Bogotá D.C.",
]

CIUDAD_A_DEPTO = {
    "bogota": "Bogotá D.C.", "medellin": "Antioquia", "envigado": "Antioquia",
    "itagui": "Antioquia", "bello": "Antioquia", "rionegro": "Antioquia",
    "cali": "Valle del Cauca", "palmira": "Valle del Cauca", "buga": "Valle del Cauca",
    "barranquilla": "Atlántico", "soledad": "Atlántico", "malambo": "Atlántico",
    "cartagena": "Bolívar", "bucaramanga": "Santander", "floridablanca": "Santander",
    "giron": "Santander", "piedecuesta": "Santander", "cucuta": "Norte de Santander",
    "pereira": "Risaralda", "dosquebradas": "Risaralda", "manizales": "Caldas",
    "armenia": "Quindío", "ibague": "Tolima", "villavicencio": "Meta",
    "santa marta": "Magdalena", "monteria": "Córdoba", "neiva": "Huila",
    "pasto": "Nariño", "popayan": "Cauca", "valledupar": "Cesar",
    "sincelejo": "Sucre", "tunja": "Boyacá", "riohacha": "La Guajira",
    "quibdo": "Chocó", "florencia": "Caquetá", "yopal": "Casanare",
    "arauca": "Arauca", "mocoa": "Putumayo", "leticia": "Amazonas",
}

NIVELES_EDUCATIVOS = [
    ("doctorado", "Doctorado"), ("phd", "Doctorado"),
    ("maestria", "Maestría"), ("magister", "Maestría"), ("mba", "Maestría"),
    ("especializacion", "Especialización"), ("especialista", "Especialización"),
    ("profesional", "Profesional"), ("pregrado", "Profesional"),
    ("universitario", "Profesional"), ("ingenier", "Profesional"),
    ("licenciad", "Profesional"), ("administrador", "Profesional"),
    ("tecnologo", "Tecnólogo"), ("tecnologia", "Tecnólogo"),
    ("tecnico", "Técnico"),
    ("bachiller", "Bachiller"), ("secundaria", "Bachiller"),
]

NIVEL_ORDEN = {
    "Bachiller": 1, "Técnico": 2, "Tecnólogo": 3, "Profesional": 4,
    "Especialización": 5, "Maestría": 6, "Doctorado": 7,
}

IDIOMAS = {
    "ingles": "Inglés", "english": "Inglés", "frances": "Francés",
    "portugues": "Portugués", "aleman": "Alemán", "italiano": "Italiano",
    "mandarin": "Mandarín", "espanol": "Español",
}

NIVEL_IDIOMA = re.compile(
    r"\b(a1|a2|b1|b2|c1|c2|basico|básico|intermedio|avanzado|nativo|bilingue|bilingüe)\b",
    re.IGNORECASE,
)

# Habilidades frecuentes en las campañas del ATS.
HABILIDADES = [
    "Excel", "Word", "PowerPoint", "SAP", "Salesforce", "HubSpot", "Zendesk",
    "CRM", "Call center", "Servicio al cliente", "Atención al cliente",
    "Ventas", "Telemercadeo", "Cobranza", "Recaudo", "Cartera", "Facturación",
    "Contabilidad", "NIIF", "Nómina", "Siigo", "World Office", "Helisa",
    "SQL", "Python", "Java", "JavaScript", "React", "Node", "Linux", "Windows Server",
    "Redes", "Soporte técnico", "Mesa de ayuda", "ITIL", "Active Directory",
    "Enfermería", "Auxiliar de enfermería", "Primeros auxilios", "Triage",
    "Historia clínica", "Facturación en salud", "Liderazgo", "Trabajo en equipo",
]

CERTIFICACIONES = re.compile(
    r"\b(scrum master|pmp|itil|cisco|ccna|comptia|aws certified|azure|"
    r"six sigma|iso 9001|sst|salud ocupacional|manipulaci[oó]n de alimentos|"
    r"alturas|reanimaci[oó]n|soporte vital)\b",
    re.IGNORECASE,
)


# ─────────────────────────── Modelo de salida ──────────────────────────────

@dataclass
class Extraccion:
    nombres: str | None = None
    apellidos: str | None = None
    cedula: str | None = None
    tel: str | None = None
    telAlt: str | None = None
    email: str | None = None
    ciudad: str | None = None
    depto: str | None = None
    nacimiento: str | None = None
    educacion: str | None = None
    universidad: str | None = None
    experiencia: float | None = None
    cargoActual: str | None = None
    habilidades: list[str] = field(default_factory=list)
    idiomas: list[str] = field(default_factory=list)
    certificaciones: list[str] = field(default_factory=list)
    # Confianza por campo: la interfaz marca en ámbar lo que conviene revisar.
    confianza: dict[str, float] = field(default_factory=dict)


# ─────────────────────────── Extracción por campo ──────────────────────────

def extraer_cedula(texto: str) -> tuple[str | None, float]:
    m = RE_CEDULA.search(texto)
    if m:
        digitos = re.sub(r"\D", "", m.group(1))
        if 6 <= len(digitos) <= 10:
            return digitos, 0.95

    # Sin etiqueta: un número suelto de 8 a 10 dígitos con puntos de miles.
    for m in re.finditer(r"\b\d{1,3}(?:\.\d{3}){2,3}\b", texto):
        digitos = re.sub(r"\D", "", m.group(0))
        if 7 <= len(digitos) <= 10:
            return digitos, 0.6
    return None, 0.0


def extraer_telefonos(texto: str) -> tuple[str | None, str | None, float]:
    celulares: list[str] = []
    for m in RE_CELULAR.finditer(texto):
        numero = f"+57 {m.group(1)} {m.group(2)} {m.group(3)}"
        if numero not in celulares:
            celulares.append(numero)

    if celulares:
        alterno = celulares[1] if len(celulares) > 1 else None
        if not alterno:
            f = RE_FIJO.search(texto)
            if f:
                alterno = f"({f.group(1)}) {f.group(2)} {f.group(3)}"
        return celulares[0], alterno, 0.9

    f = RE_FIJO.search(texto)
    if f:
        return f"({f.group(1)}) {f.group(2)} {f.group(3)}", None, 0.7
    return None, None, 0.0


def extraer_email(texto: str) -> tuple[str | None, float]:
    correos = RE_EMAIL.findall(texto)
    if not correos:
        return None, 0.0
    # El personal suele ir primero; se descartan los de la empresa anterior.
    for correo in correos:
        if not re.search(r"@(empresa|company|corp)\.", correo, re.IGNORECASE):
            return correo.lower(), 0.95
    return correos[0].lower(), 0.8


def extraer_nombre(texto: str) -> tuple[str | None, str | None, float]:
    lineas = [l.strip() for l in texto.split("\n") if l.strip()][:12]

    etiquetado = re.search(
        r"(?:nombre[s]?\s*(?:completo)?)\s*[:\-]\s*([A-ZÁÉÍÓÚÑ][^\n]{4,60})",
        texto, re.IGNORECASE,
    )
    candidatas = ([etiquetado.group(1)] if etiquetado else []) + lineas

    for linea in candidatas:
        limpia = re.sub(r"[^\wÁÉÍÓÚÑáéíóúñ\s]", " ", linea).strip()
        limpia = re.sub(r"\s+", " ", limpia)
        palabras = limpia.split()

        if not (2 <= len(palabras) <= 5):
            continue
        if any(ch.isdigit() for ch in limpia):
            continue
        # Descarta encabezados típicos
        if re.search(
            r"\b(hoja|vida|curriculum|currículum|resume|perfil|datos|personales|"
            r"profesional|experiencia|contacto)\b",
            sin_tildes(limpia), re.IGNORECASE,
        ):
            continue
        if not all(p[0].isupper() for p in palabras if len(p) > 2):
            continue

        confianza = 0.9 if etiquetado and linea == etiquetado.group(1) else 0.7
        if len(palabras) >= 4:
            corte = len(palabras) // 2
        elif len(palabras) == 3:
            corte = 1
        else:
            corte = 1
        return " ".join(palabras[:corte]), " ".join(palabras[corte:]), confianza

    return None, None, 0.0


def extraer_ubicacion(texto: str) -> tuple[str | None, str | None, float]:
    plano = sin_tildes(texto).lower()

    ciudad = depto = None
    for clave, departamento in CIUDAD_A_DEPTO.items():
        if re.search(rf"\b{re.escape(clave)}\b", plano):
            ciudad = clave.title() if clave != "bogota" else "Bogotá D.C."
            depto = departamento
            break

    if not depto:
        for departamento in DEPARTAMENTOS:
            if re.search(rf"\b{re.escape(sin_tildes(departamento).lower())}\b", plano):
                depto = departamento
                break

    return ciudad, depto, 0.8 if ciudad else (0.6 if depto else 0.0)


def extraer_nacimiento(texto: str) -> tuple[str | None, float]:
    contexto = re.search(
        r"(?:nacimiento|naci[oó]|fecha de nac)[^\n]{0,60}", texto, re.IGNORECASE
    )
    fuente = contexto.group(0) if contexto else texto[:1200]

    for m in RE_FECHA.finditer(fuente):
        dia, mes_txt, anio = m.group(1), m.group(2).lower()[:3], int(m.group(3))
        mes = MESES.get(mes_txt) or (int(mes_txt) if mes_txt.isdigit() else None)
        if not mes or not (1 <= mes <= 12):
            continue
        # Rango razonable para un candidato laboral
        if not (date.today().year - 75 <= anio <= date.today().year - 16):
            continue
        try:
            return date(anio, mes, int(dia)).isoformat(), 0.85 if contexto else 0.5
        except ValueError:
            continue
    return None, 0.0


def extraer_educacion(texto: str) -> tuple[str | None, str | None, float]:
    plano = sin_tildes(texto).lower()

    nivel = None
    for clave, etiqueta in NIVELES_EDUCATIVOS:
        if clave in plano:
            if nivel is None or NIVEL_ORDEN[etiqueta] > NIVEL_ORDEN[nivel]:
                nivel = etiqueta

    universidad = None
    m = re.search(
        r"\b((?:universidad|corporaci[oó]n|fundaci[oó]n|instituci[oó]n|"
        r"polit[eé]cnico|sena|colegio mayor|escuela)[^\n,;.]{3,60})",
        texto, re.IGNORECASE,
    )
    if m:
        universidad = re.sub(r"\s+", " ", m.group(1)).strip().title()

    return nivel, universidad, 0.8 if nivel else 0.0


def extraer_experiencia(texto: str) -> tuple[float | None, float]:
    m = re.search(
        r"(\d{1,2})(?:[.,](\d))?\s*(?:\+)?\s*a[ñn]os?\s+(?:de\s+)?experiencia",
        texto, re.IGNORECASE,
    )
    if m:
        valor = float(f"{m.group(1)}.{m.group(2) or 0}")
        if 0 < valor <= 50:
            return valor, 0.9

    # Sin declaración explícita: se estima por los años de los empleos.
    anios = sorted({int(a) for a in re.findall(r"\b(19[89]\d|20[0-3]\d)\b", texto)})
    if len(anios) >= 2:
        estimado = min(date.today().year, anios[-1]) - anios[0]
        if 0 < estimado <= 50:
            return float(estimado), 0.45
    return None, 0.0


def extraer_cargo(texto: str) -> tuple[str | None, float]:
    m = re.search(
        r"(?:cargo\s*actual|puesto\s*actual|cargo)\s*[:\-]\s*([^\n]{3,60})",
        texto, re.IGNORECASE,
    )
    if m:
        return re.sub(r"\s+", " ", m.group(1)).strip(), 0.85

    m = re.search(
        r"\b(actualmente|actual)\b[^\n]{0,20}\b(?:como|de)\s+([^\n,.;]{3,50})",
        texto, re.IGNORECASE,
    )
    if m:
        return m.group(2).strip().title(), 0.6
    return None, 0.0


def extraer_habilidades(texto: str) -> list[str]:
    plano = sin_tildes(texto).lower()
    halladas = [
        h for h in HABILIDADES
        if re.search(rf"\b{re.escape(sin_tildes(h).lower())}\b", plano)
    ]
    return halladas[:20]


def extraer_idiomas(texto: str) -> list[str]:
    plano = sin_tildes(texto).lower()
    salida: list[str] = []
    for clave, nombre in IDIOMAS.items():
        pos = plano.find(clave)
        if pos < 0:
            continue
        cerca = texto[max(0, pos - 40): pos + 80]
        nivel = NIVEL_IDIOMA.search(cerca)
        salida.append(f"{nombre} {nivel.group(1).upper()}" if nivel else nombre)
    return salida


def extraer_certificaciones(texto: str) -> list[str]:
    vistas: list[str] = []
    for m in CERTIFICACIONES.finditer(texto):
        valor = m.group(1).title()
        if valor not in vistas:
            vistas.append(valor)
    return vistas[:10]


# ─────────────────────────── Orquestación ──────────────────────────────────

def analizar(texto: str) -> Extraccion:
    r = Extraccion()

    r.nombres, r.apellidos, c_nombre = extraer_nombre(texto)
    r.cedula, c_cedula = extraer_cedula(texto)
    r.tel, r.telAlt, c_tel = extraer_telefonos(texto)
    r.email, c_email = extraer_email(texto)
    r.ciudad, r.depto, c_ubic = extraer_ubicacion(texto)
    r.nacimiento, c_nac = extraer_nacimiento(texto)
    r.educacion, r.universidad, c_edu = extraer_educacion(texto)
    r.experiencia, c_exp = extraer_experiencia(texto)
    r.cargoActual, c_cargo = extraer_cargo(texto)

    r.habilidades = extraer_habilidades(texto)
    r.idiomas = extraer_idiomas(texto)
    r.certificaciones = extraer_certificaciones(texto)

    r.confianza = {
        "nombre": c_nombre, "cedula": c_cedula, "tel": c_tel, "email": c_email,
        "ubicacion": c_ubic, "nacimiento": c_nac, "educacion": c_edu,
        "experiencia": c_exp, "cargoActual": c_cargo,
        "habilidades": 0.75 if r.habilidades else 0.0,
        "idiomas": 0.8 if r.idiomas else 0.0,
    }
    return r


# ─────────────────────────── API ───────────────────────────────────────────

@app.get("/health")
def salud():
    return {"ok": True, "servicio": "extractor", "version": "1.0.0"}


@app.post("/extract")
async def extraer(archivo: UploadFile = File(...)):
    """
    Recibe la hoja de vida y devuelve los campos extraídos con su nivel de
    confianza. No guarda nada: los bytes se descartan al terminar.
    """
    datos = await archivo.read()

    if not datos:
        raise HTTPException(400, "El archivo está vacío")
    if len(datos) > MAX_BYTES:
        raise HTTPException(413, "El archivo supera los 10 MB")

    texto = extraer_texto(datos, archivo.filename or "")

    if len(texto) < 60:
        return JSONResponse(
            status_code=422,
            content={
                "error": "No se pudo leer texto del archivo. "
                         "Si es un PDF escaneado, necesita OCR.",
                "codigo": "sin_texto",
                "caracteres": len(texto),
            },
        )

    resultado = analizar(texto)
    salida = asdict(resultado)
    salida["meta"] = {
        "caracteres": len(texto),
        "archivo": archivo.filename,
        "camposDetectados": sum(
            1 for k, v in salida.items()
            if k not in ("confianza", "meta") and v not in (None, [], "")
        ),
    }
    return salida
